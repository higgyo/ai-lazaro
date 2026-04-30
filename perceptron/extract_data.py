import docx
import pandas as pd
import io
import re

def extract_tables_to_csv(doc_path):
    doc = docx.Document(doc_path)
    
    # In python-docx, we can access doc.tables
    tables = doc.tables
    
    # We expect 3 tables in the doc based on its structure:
    # 1. Registration for weights and epochs
    # 2. Registration for output classes (test set)
    # 3. Training dataset (in the annex)
    
    # We want the training dataset and the test samples.
    # The training dataset is likely the 3rd table.
    # Let's search for the tables by header
    
    treinamento_data = []
    teste_data = []
    
    for t in tables:
        # Check first row
        first_row = [cell.text.strip() for cell in t.rows[0].cells]
        
        # Test table typically has columns: Amostra, x1, x2, x3, y(T1)...
        if 'Amostra' in first_row and 'x1' in ''.join(first_row).replace(' ', ''):
            # Extract test
            # Skip the header row
            for row in t.rows[1:]:
                cells = [c.text.strip().replace(',', '.') for c in row.cells]
                # the sample values are Amostra, x1, x2, x3
                # the rest are empty spaces to be filled
                if len(cells) >= 4:
                    try:
                        x1 = float(cells[1])
                        x2 = float(cells[2])
                        x3 = float(cells[3])
                        teste_data.append([cells[0], x1, x2, x3])
                    except ValueError:
                        continue
                        
        elif 'Padrão' in first_row or 'x1' in ''.join(first_row).replace(' ', ''):
            # Check if it has 'd' (desired output)
            if 'd' in first_row:
                for row in t.rows[1:]:
                    cells = [c.text.strip().replace(',', '.') for c in row.cells]
                    if len(cells) >= 5:
                        try:
                            # cells[0] is Padrão
                            x1 = float(cells[1])
                            x2 = float(cells[2])
                            x3 = float(cells[3])
                            d = float(cells[4])
                            treinamento_data.append([cells[0], x1, x2, x3, d])
                        except ValueError:
                            continue

    df_treino = pd.DataFrame(treinamento_data, columns=['Padrao', 'x1', 'x2', 'x3', 'd'])
    df_treino.to_csv('treinamento.csv', index=False)
    
    df_teste = pd.DataFrame(teste_data, columns=['Amostra', 'x1', 'x2', 'x3'])
    df_teste.to_csv('teste.csv', index=False)

if __name__ == "__main__":
    extract_tables_to_csv('Perceptron.docx')
    print("CSV files generated.")
