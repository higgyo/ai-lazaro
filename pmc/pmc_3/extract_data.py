import docx
import csv
import re

doc = docx.Document('PMC3.docx')

data = {}

# Extract from Table 3 (Training)
table3 = doc.tables[3]
for row in table3.rows[1:]:
    cells = [cell.text.strip() for cell in row.cells]
    # pairs are 0,1; 2,3; 4,5; 6,7
    for i in range(0, 8, 2):
        t_str = cells[i]
        f_str = cells[i+1]
        if t_str.startswith('t = '):
            t = int(t_str.replace('t = ', ''))
            f = float(f_str.replace(',', '.'))
            data[t] = f

# Extract from Table 2 (Validation)
table2 = doc.tables[2]
for row in table2.rows[2:]:
    cells = [cell.text.strip() for cell in row.cells]
    t_str = cells[0]
    f_str = cells[1]
    if t_str.startswith('t = '):
        t = int(t_str.replace('t = ', ''))
        f = float(f_str.replace(',', '.'))
        data[t] = f

# Check if all t from 1 to 120 are present
missing = [t for t in range(1, 121) if t not in data]
if missing:
    print(f"Missing data for t: {missing}")
else:
    print("All 120 data points successfully extracted.")

# Write to serie_temporal.csv
with open('serie_temporal.csv', 'w', newline='') as f:
    writer = csv.writer(f)
    writer.writerow(['t', 'f(t)'])
    for t in range(1, 121):
        if t in data:
            writer.writerow([t, data[t]])

print("Saved to serie_temporal.csv")
